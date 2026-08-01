"""
Aplicatie evidenta turisti - punct de intrare Flask.

Ruleaza cu:  python app.py
Apoi deschide in browser:  http://localhost:5000
De pe alt PC din aceeasi retea:  http://<IP-ul-acestui-calculator>:5000
"""
import contextlib
import io
import os
import uuid
from datetime import datetime, date

from flask import (
    Flask, render_template, request, redirect, url_for, flash,
    send_from_directory, send_file, abort,
)
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from sqlalchemy import inspect, text as sql_text

from models import db, Turist, Excursie, Inscriere, CostExcursie
from alerts import (
    atentionari_excursie, turisti_cu_premiu_de_acordat,
    zile_speciale_apropiate, gaseste_onomastica, formateaza_zi_luna,
)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
INSTANCE_DIR = os.path.join(BASE_DIR, "instance")
UPLOAD_DIR = os.path.join(INSTANCE_DIR, "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)

EXTENSII_PERMISE_POZA = {"png", "jpg", "jpeg", "webp", "pdf"}

app = Flask(__name__)
app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///" + os.path.join(INSTANCE_DIR, "evidenta.db")
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024  # 16 MB per poza, suficient
app.secret_key = "schimba-aceasta-cheie-daca-pui-aplicatia-pe-internet"

db.init_app(app)


def asigura_schema_actualizata():
    """
    Creeaza tabelele care lipsesc si adauga coloanele noi aparute intre
    versiunile aplicatiei, FARA sa stearga vreo data existenta.

    SQLite nu actualizeaza singur structura unui tabel deja creat cand
    modificam models.py (ex: cand am adaugat 'poza_ci') - de-asta se rula
    fix aici, automat, de fiecare data cand aplicatia porneste.
    """
    db.create_all()
    inspector = inspect(db.engine)
    for model in (Turist, Excursie, Inscriere, CostExcursie):
        tabel = model.__tablename__
        if tabel not in inspector.get_table_names():
            continue
        coloane_existente = {c["name"] for c in inspector.get_columns(tabel)}
        for coloana in model.__table__.columns:
            if coloana.name in coloane_existente:
                continue
            tip_sql = coloana.type.compile(dialect=db.engine.dialect)
            print(f"[migrare bd] adaug coloana lipsa: {tabel}.{coloana.name} ({tip_sql})")
            with db.engine.begin() as conn:
                conn.execute(sql_text(f'ALTER TABLE {tabel} ADD COLUMN "{coloana.name}" {tip_sql}'))


with app.app_context():
    asigura_schema_actualizata()


def extensie_permisa(nume_fisier):
    return "." in nume_fisier and nume_fisier.rsplit(".", 1)[1].lower() in EXTENSII_PERMISE_POZA


def salveaza_poza_ci(fisier):
    """Salveaza fisierul uploadat cu un nume unic si returneaza numele salvat."""
    if not fisier or not fisier.filename or not extensie_permisa(fisier.filename):
        return None
    ext = fisier.filename.rsplit(".", 1)[1].lower()
    nume_unic = f"{uuid.uuid4().hex}.{ext}"
    fisier.save(os.path.join(UPLOAD_DIR, secure_filename(nume_unic)))
    return nume_unic


def sterge_poza_ci(nume_fisier):
    if not nume_fisier:
        return
    cale = os.path.join(UPLOAD_DIR, nume_fisier)
    if os.path.exists(cale):
        os.remove(cale)


@app.route("/uploads/<path:nume_fisier>")
def fisier_incarcat(nume_fisier):
    return send_from_directory(UPLOAD_DIR, nume_fisier)


def parse_data(text):
    """Converteste un text din formular (YYYY-MM-DD) intr-un obiect date."""
    if not text:
        return None
    try:
        return datetime.strptime(text, "%Y-%m-%d").date()
    except ValueError:
        return None


# --------------------------------------------------------------------------
# Pagina principala
# --------------------------------------------------------------------------
@app.route("/")
def index():
    nr_turisti = Turist.query.count()
    nr_excursii = Excursie.query.count()
    ultimele_excursii = (
        Excursie.query.order_by(Excursie.data_inceput.desc()).limit(5).all()
    )
    turisti_premiu = turisti_cu_premiu_de_acordat(Turist.query.all())

    extins = request.args.get("curand") == "extins"
    luni = request.args.get("luni", "2")
    luni = int(luni) if luni.isdigit() else 2

    toate_zilele = zile_speciale_apropiate(Turist.query.all())
    if extins:
        zile_curand = [z for z in toate_zilele if z["zile_pana_la"] <= luni * 30]
    else:
        zile_curand = toate_zilele[:3]

    return render_template(
        "index.html",
        nr_turisti=nr_turisti,
        nr_excursii=nr_excursii,
        ultimele_excursii=ultimele_excursii,
        turisti_premiu=turisti_premiu,
        zile_curand=zile_curand,
        extins=extins,
        luni=luni,
    )


# --------------------------------------------------------------------------
# Turisti
# --------------------------------------------------------------------------
@app.route("/turisti")
def turisti_list():
    q = request.args.get("q", "").strip()
    query = Turist.query
    if q:
        query = query.filter(Turist.nume.ilike(f"%{q}%"))
    turisti = query.order_by(Turist.nume).all()
    return render_template("turisti_list.html", turisti=turisti, q=q)


@app.route("/turisti/nou", methods=["GET", "POST"])
def turist_nou():
    if request.method == "POST":
        t = Turist(nume=request.form["nume"])
        _completeaza_turist(t, request.form)
        t.poza_ci = salveaza_poza_ci(request.files.get("poza_ci"))
        db.session.add(t)
        db.session.commit()
        flash(f"Turistul {t.nume} a fost adaugat.", "success")
        return redirect(url_for("turisti_list"))
    return render_template("turist_form.html", turist=None)


@app.route("/turisti/<int:turist_id>")
def turist_detaliu(turist_id):
    turist = Turist.query.get_or_404(turist_id)
    onomastica = gaseste_onomastica(turist.nume)
    onomastica_text = formateaza_zi_luna(*onomastica) if onomastica else None
    return render_template("turist_detaliu.html", turist=turist, onomastica_text=onomastica_text)


@app.route("/turisti/<int:turist_id>/editeaza", methods=["GET", "POST"])
def turist_editeaza(turist_id):
    turist = Turist.query.get_or_404(turist_id)
    if request.method == "POST":
        schimbari = []
        if turist.nume != request.form["nume"]:
            schimbari.append(f"nume: '{turist.nume}' -> '{request.form['nume']}'")
        if turist.adresa != request.form.get("adresa"):
            schimbari.append(f"adresa: '{turist.adresa}' -> '{request.form.get('adresa')}'")
        if turist.serie_ci != request.form.get("serie_ci"):
            schimbari.append(f"CI: '{turist.serie_ci}' -> '{request.form.get('serie_ci')}'")

        turist.nume = request.form["nume"]
        _completeaza_turist(turist, request.form)

        if request.form.get("elimina_poza"):
            sterge_poza_ci(turist.poza_ci)
            turist.poza_ci = None
        else:
            poza_noua = salveaza_poza_ci(request.files.get("poza_ci"))
            if poza_noua:
                sterge_poza_ci(turist.poza_ci)
                turist.poza_ci = poza_noua

        if schimbari:
            linie = f"[{date.today().isoformat()}] " + "; ".join(schimbari)
            turist.istoric_modificari = (
                (turist.istoric_modificari + "\n" + linie) if turist.istoric_modificari else linie
            )

        db.session.commit()
        flash("Datele turistului au fost actualizate.", "success")
        return redirect(url_for("turist_detaliu", turist_id=turist.id))
    return render_template("turist_form.html", turist=turist)


@app.route("/turisti/<int:turist_id>/extra", methods=["POST"])
def turist_extra(turist_id):
    turist = Turist.query.get_or_404(turist_id)
    turist.notite_extra = request.form.get("notite_extra") or None
    db.session.commit()
    flash("Notitele extra au fost salvate.", "success")
    return redirect(url_for("turist_detaliu", turist_id=turist.id))


@app.route("/turisti/<int:turist_id>/sterge", methods=["POST"])
def turist_sterge(turist_id):
    turist = Turist.query.get_or_404(turist_id)
    sterge_poza_ci(turist.poza_ci)
    db.session.delete(turist)
    db.session.commit()
    flash("Turist sters.", "info")
    return redirect(url_for("turisti_list"))


def clean(value):
    if value is None:
        return None

    value = value.strip()

    if value == "" or value.lower() == "none":
        return None

    return value


def _int_sigur(text):
    text = (text or "").strip()
    if not text:
        return 0
    try:
        return int(text)
    except ValueError:
        return 0


def _completeaza_turist(t, form):
    t.cnp = clean(form.get("cnp"))
    t.serie_ci = clean(form.get("serie_ci"))
    t.emisa_de = clean(form.get("emisa_de"))
    t.adresa = clean(form.get("adresa"))
    t.localitate = clean(form.get("localitate"))
    t.judet = clean(form.get("judet"))
    t.data_nasterii = parse_data(form.get("data_nasterii"))
    t.telefon = clean(form.get("telefon"))
    t.email = clean(form.get("email"))
    t.observatii = clean(form.get("observatii"))
    t.ajustare_excursii = _int_sigur(form.get("ajustare_excursii"))
    t.ajustare_cadouri = _int_sigur(form.get("ajustare_cadouri"))


# --------------------------------------------------------------------------
# Excursii
# --------------------------------------------------------------------------
@app.route("/excursii")
def excursii_list():
    q = request.args.get("q", "").strip()
    an_selectat = request.args.get("an", "").strip()
    ordine = request.args.get("ordine", "desc")

    query = Excursie.query
    if q:
        query = query.filter(Excursie.nume.ilike(f"%{q}%"))
    if an_selectat.isdigit():
        query = query.filter(db.extract("year", Excursie.data_inceput) == int(an_selectat))

    coloana_data = Excursie.data_inceput
    query = query.order_by(coloana_data.asc() if ordine == "asc" else coloana_data.desc())
    excursii = query.all()

    ani_disponibili = sorted(
        {e.data_inceput.year for e in Excursie.query.all() if e.data_inceput}, reverse=True
    )

    return render_template(
        "excursii_list.html",
        excursii=excursii, q=q,
        an_selectat=an_selectat, ordine=ordine, ani_disponibili=ani_disponibili,
    )


@app.route("/excursii/noua", methods=["GET", "POST"])
def excursie_noua():
    if request.method == "POST":
        e = Excursie(nume=request.form["nume"])
        _completeaza_excursie(e, request.form)
        db.session.add(e)
        db.session.commit()
        flash(f"Excursia {e.nume} a fost adaugata.", "success")
        return redirect(url_for("excursii_list"))
    return render_template("excursie_form.html", excursie=None)


@app.route("/excursii/<int:excursie_id>")
def excursie_detaliu(excursie_id):
    excursie = Excursie.query.get_or_404(excursie_id)
    turisti_inscrisi_ids = {i.turist_id for i in excursie.inscrieri}
    turisti_disponibili = (
        Turist.query.filter(~Turist.id.in_(turisti_inscrisi_ids))
        .order_by(Turist.nume)
        .all()
        if turisti_inscrisi_ids
        else Turist.query.order_by(Turist.nume).all()
    )
    atentionari = atentionari_excursie(excursie)
    inscrieri_de_achitat = [i for i in excursie.inscrieri if not i.excursie_cadou]
    return render_template(
        "excursie_detaliu.html",
        excursie=excursie,
        turisti_disponibili=turisti_disponibili,
        atentionari=atentionari,
        inscrieri_de_achitat=inscrieri_de_achitat,
    )


@app.route("/excursii/<int:excursie_id>/editeaza", methods=["GET", "POST"])
def excursie_editeaza(excursie_id):
    excursie = Excursie.query.get_or_404(excursie_id)
    if request.method == "POST":
        excursie.nume = request.form["nume"]
        _completeaza_excursie(excursie, request.form)
        db.session.commit()
        flash("Excursia a fost actualizata.", "success")
        return redirect(url_for("excursie_detaliu", excursie_id=excursie.id))
    return render_template("excursie_form.html", excursie=excursie)


@app.route("/excursii/<int:excursie_id>/extra", methods=["POST"])
def excursie_extra(excursie_id):
    excursie = Excursie.query.get_or_404(excursie_id)
    excursie.notite_extra = request.form.get("notite_extra") or None
    db.session.commit()
    flash("Notitele extra au fost salvate.", "success")
    return redirect(url_for("excursie_detaliu", excursie_id=excursie.id))


@app.route("/excursii/<int:excursie_id>/sterge", methods=["POST"])
def excursie_sterge(excursie_id):
    excursie = Excursie.query.get_or_404(excursie_id)
    db.session.delete(excursie)
    db.session.commit()
    flash("Excursie stearsa.", "info")
    return redirect(url_for("excursii_list"))


@app.route("/excursii/<int:excursie_id>/autocar")
def excursie_autocar_docx(excursie_id):
    excursie = Excursie.query.get_or_404(excursie_id)
    if not excursie.inscrieri:
        flash("Nu exista turisti inscrisi la aceasta excursie inca.", "warning")
        return redirect(url_for("excursie_detaliu", excursie_id=excursie.id))

    buffer = genereaza_docx_autocar(excursie)
    nume_fisier = f"autocar_{(excursie.cod or excursie.id)}.docx"
    return send_file(
        buffer,
        as_attachment=True,
        download_name=nume_fisier,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def genereaza_docx_autocar(excursie):
    doc = Document()

    # margini mai mici, ca tabelul sa se intinda aproape de ambele margini
    for sectiune in doc.sections:
        sectiune.left_margin = Cm(1.3)
        sectiune.right_margin = Cm(1.3)
        sectiune.top_margin = Cm(1.5)
        sectiune.bottom_margin = Cm(1.5)
    latime_utila = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin

    doc.add_heading(f"Lista autocar - {excursie.nume}", level=1)

    perioada = ""
    if excursie.data_inceput:
        perioada = excursie.data_inceput.strftime("%d.%m.%Y")
        if excursie.data_sfarsit and excursie.data_sfarsit != excursie.data_inceput:
            perioada += " - " + excursie.data_sfarsit.strftime("%d.%m.%Y")
    sub = doc.add_paragraph()
    sub.add_run(f"Perioada: {perioada or '-'}    |    Total turisti: {excursie.numar_participanti}").italic = True

    if excursie.punct_contact:
        doc.add_paragraph(f"Punct de plecare / contact: {excursie.punct_contact}")
    if excursie.transport:
        doc.add_paragraph(f"Transport: {excursie.transport}")
    if excursie.ghid:
        doc.add_paragraph(f"Ghid: {excursie.ghid}")

    doc.add_paragraph()

    # latimi coloane in cm: Nr. ingust, Nume lat, restul impartite pe potriva
    latimi_cm = [1.0, 6.2, 3.2, 3.4, 4.2]
    antete = ["Nr.", "Nume turist", "Telefon", "Observatii", "Semnatura"]

    tabel = doc.add_table(rows=1, cols=len(antete))
    tabel.style = "Table Grid"
    _seteaza_latimi_coloane(tabel, latimi_cm)

    antet_celule = tabel.rows[0].cells
    for cell, text_antet in zip(antet_celule, antete):
        cell.text = text_antet
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    inscrieri_sortate = sorted(excursie.inscrieri, key=lambda i: i.turist.nume)
    for idx, inscriere in enumerate(inscrieri_sortate, start=1):
        rand = tabel.add_row().cells
        rand[0].text = str(idx)
        rand[1].text = inscriere.turist.nume
        rand[2].text = inscriere.turist.telefon or ""
        rand[3].text = inscriere.observatii or ""
        rand[4].text = ""  # semnatura - ramane goala, se completeaza pe hartie
    _seteaza_latimi_coloane(tabel, latimi_cm)  # Word ignora uneori latimea daca nu e reaplicata dupa randuri

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


@app.route("/excursii/<int:excursie_id>/lista-contact")
def excursie_lista_contact_docx(excursie_id):
    excursie = Excursie.query.get_or_404(excursie_id)
    if not excursie.inscrieri:
        flash("Nu exista turisti inscrisi la aceasta excursie inca.", "warning")
        return redirect(url_for("excursie_detaliu", excursie_id=excursie.id))

    buffer = genereaza_docx_contact(excursie)
    nume_fisier = f"contact_{(excursie.cod or excursie.id)}.docx"
    return send_file(
        buffer,
        as_attachment=True,
        download_name=nume_fisier,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def genereaza_docx_contact(excursie):
    """Export separat de lista de autocar - centrat pe date de contact
    (telefon + email), nu pe semnaturi."""
    doc = Document()
    for sectiune in doc.sections:
        sectiune.left_margin = Cm(1.3)
        sectiune.right_margin = Cm(1.3)
        sectiune.top_margin = Cm(1.5)
        sectiune.bottom_margin = Cm(1.5)

    doc.add_heading(f"Lista contact - {excursie.nume}", level=1)

    perioada = ""
    if excursie.data_inceput:
        perioada = excursie.data_inceput.strftime("%d.%m.%Y")
        if excursie.data_sfarsit and excursie.data_sfarsit != excursie.data_inceput:
            perioada += " - " + excursie.data_sfarsit.strftime("%d.%m.%Y")
    sub = doc.add_paragraph()
    sub.add_run(f"Perioada: {perioada or '-'}    |    Total turisti: {excursie.numar_participanti}").italic = True
    doc.add_paragraph()

    latimi_cm = [1.0, 6.0, 3.5, 6.0]
    antete = ["Nr.", "Nume turist", "Telefon", "E-mail"]

    tabel = doc.add_table(rows=1, cols=len(antete))
    tabel.style = "Table Grid"
    _seteaza_latimi_coloane(tabel, latimi_cm)

    antet_celule = tabel.rows[0].cells
    for cell, text_antet in zip(antet_celule, antete):
        cell.text = text_antet
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    inscrieri_sortate = sorted(excursie.inscrieri, key=lambda i: i.turist.nume)
    for idx, inscriere in enumerate(inscrieri_sortate, start=1):
        rand = tabel.add_row().cells
        rand[0].text = str(idx)
        rand[1].text = inscriere.turist.nume
        rand[2].text = inscriere.turist.telefon or ""
        rand[3].text = inscriere.turist.email or ""
    _seteaza_latimi_coloane(tabel, latimi_cm)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _seteaza_latimi_coloane(tabel, latimi_cm):
    """
    python-docx / Word ignora frecvent table.columns[i].width daca tabelul
    nu e marcat explicit cu layout 'fix' si daca latimea nu e setata si pe
    fiecare celula in parte - de-aia facem ambele mai jos.
    """
    tabel.autofit = False
    tbl_pr = tabel._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)

    for rand in tabel.rows:
        for idx, latime in enumerate(latimi_cm):
            rand.cells[idx].width = Cm(latime)
    for idx, latime in enumerate(latimi_cm):
        tabel.columns[idx].width = Cm(latime)


def _completeaza_excursie(e, form):
    e.cod = form.get("cod") or None
    e.data_inceput = parse_data(form.get("data_inceput"))
    e.data_sfarsit = parse_data(form.get("data_sfarsit"))
    nr_zile = form.get("nr_zile")
    e.nr_zile = int(nr_zile) if nr_zile else None
    e.obiective = form.get("obiective")
    e.cazare = form.get("cazare")
    e.transport = form.get("transport")
    e.ghid = form.get("ghid")
    e.punct_contact = form.get("punct_contact")
    e.observatii = form.get("observatii")
    pret = form.get("pret")
    e.pret = float(pret) if pret else None


# --------------------------------------------------------------------------
# Inscrieri (leaga un turist de o excursie)
# --------------------------------------------------------------------------
@app.route("/inscrieri/noua", methods=["POST"])
def inscriere_noua():
    excursie_id = int(request.form["excursie_id"])
    turist_id = int(request.form["turist_id"])

    exista = Inscriere.query.filter_by(
        excursie_id=excursie_id, turist_id=turist_id
    ).first()
    if exista:
        flash("Acest turist este deja inscris la aceasta excursie.", "warning")
        return redirect(url_for("excursie_detaliu", excursie_id=excursie_id))

    i = Inscriere(
        turist_id=turist_id,
        excursie_id=excursie_id,
        suma_achitata=float(request.form.get("suma_achitata") or 0),
        discount=float(request.form.get("discount") or 0),
        gratuitate=bool(request.form.get("gratuitate")),
        penalizare=float(request.form.get("penalizare") or 0),
        observatii=request.form.get("observatii"),
    )
    db.session.add(i)
    db.session.commit()
    flash("Turist inscris la excursie.", "success")
    return redirect(url_for("excursie_detaliu", excursie_id=excursie_id))


@app.route("/inscrieri/<int:inscriere_id>/sterge", methods=["POST"])
def inscriere_sterge(inscriere_id):
    i = Inscriere.query.get_or_404(inscriere_id)
    excursie_id = i.excursie_id
    db.session.delete(i)
    db.session.commit()
    flash("Inscriere stearsa.", "info")
    return redirect(url_for("excursie_detaliu", excursie_id=excursie_id))


@app.route("/inscrieri/<int:inscriere_id>/toggle-cadou", methods=["POST"])
def inscriere_toggle_cadou(inscriere_id):
    i = Inscriere.query.get_or_404(inscriere_id)
    i.excursie_cadou = not i.excursie_cadou
    db.session.commit()
    return redirect(url_for("excursie_detaliu", excursie_id=i.excursie_id) + "#costuri")


@app.route("/inscrieri/<int:inscriere_id>/plata", methods=["POST"])
def inscriere_plata(inscriere_id):
    i = Inscriere.query.get_or_404(inscriere_id)
    i.suma_achitata = float(request.form.get("suma_achitata") or 0)
    if request.form.get("marcheaza_achitat"):
        i.achitat_manual = True
    db.session.commit()
    return redirect(url_for("excursie_detaliu", excursie_id=i.excursie_id) + "#costuri")


# --------------------------------------------------------------------------
# Costuri excursie (hoteluri / obiective / extra) - nu apar in Word
# --------------------------------------------------------------------------
@app.route("/excursii/<int:excursie_id>/costuri/adauga", methods=["POST"])
def cost_adauga(excursie_id):
    excursie = Excursie.query.get_or_404(excursie_id)
    tip = request.form.get("tip")
    if tip not in ("hotel", "obiectiv", "extra"):
        abort(400)

    pret = request.form.get("pret")
    cost = CostExcursie(
        excursie_id=excursie.id,
        tip=tip,
        nume=request.form.get("nume") or None,
        pret=float(pret) if pret else None,
        unitate_pret=request.form.get("unitate_pret") or None,
        contact=request.form.get("contact") or None,
        notite=request.form.get("notite") or None,
    )
    db.session.add(cost)
    db.session.commit()
    flash("Cost adaugat.", "success")
    return redirect(url_for("excursie_detaliu", excursie_id=excursie.id) + "#costuri")


@app.route("/costuri/<int:cost_id>/sterge", methods=["POST"])
def cost_sterge(cost_id):
    cost = CostExcursie.query.get_or_404(cost_id)
    excursie_id = cost.excursie_id
    db.session.delete(cost)
    db.session.commit()
    return redirect(url_for("excursie_detaliu", excursie_id=excursie_id) + "#costuri")


# --------------------------------------------------------------------------
# Import fisiere direct din aplicatie (Excel/Word), fara linia de comanda
# --------------------------------------------------------------------------
EXTENSII_PERMISE_IMPORT = {"xls", "xlsx", "doc", "docx"}


@app.route("/import", methods=["GET", "POST"])
def import_fisiere():
    if request.method == "POST":
        fisiere_incarcate = [f for f in request.files.getlist("fisiere") if f and f.filename]

        if not fisiere_incarcate:
            flash("Nu ai selectat niciun fisier.", "warning")
            return redirect(url_for("import_fisiere"))

        cai_temporare = []
        respinse = []
        avertismente = []
        for f in fisiere_incarcate:
            ext = f.filename.rsplit(".", 1)[-1].lower() if "." in f.filename else ""
            if ext not in EXTENSII_PERMISE_IMPORT:
                respinse.append(f.filename)
                continue
            nume_sigur = secure_filename(f.filename)
            cale = os.path.join(UPLOAD_DIR, f"import_tmp_{uuid.uuid4().hex}_{nume_sigur}")
            f.save(cale)

            # verificare de siguranta: uneori, mai ales pe un folder de retea,
            # scrierea fisierului nu se termina instant - verificam ca a ajuns
            # ceva pe disc inainte sa incercam sa-l citim
            marime = os.path.getsize(cale) if os.path.exists(cale) else 0
            if marime == 0:
                avertismente.append(f"'{f.filename}' s-a salvat gol (0 KB) - incearca din nou sau verifica fisierul original.")
                continue

            cai_temporare.append(cale)

        # importul se face aici, in acelasi proces - fara terminal, fara
        # VS Code. Import "lazy" (in interiorul functiei) ca sa evitam o
        # referinta circulara intre app.py si import_excel.py (acesta din
        # urma face "from app import app" la randul lui).
        import import_excel

        jurnal = io.StringIO()
        esuate = []
        try:
            with contextlib.redirect_stdout(jurnal):
                import_excel.importa_fisiere(cai_temporare)
        except Exception as exc:
            jurnal.write(f"\nEROARE neasteptata: {exc}\n")
        finally:
            for cale in cai_temporare:
                # daca fisierul respectiv a aparut ca "n-am putut deschide"
                # in jurnal, il pastram pe disc (nu il stergem), ca sa poata
                # fi verificat manual - restul, care s-au procesat cu succes,
                # se sterg normal
                if "Nu am putut deschide" in jurnal.getvalue() and os.path.basename(cale) in jurnal.getvalue():
                    esuate.append(cale)
                    continue
                try:
                    os.remove(cale)
                except OSError:
                    pass

        raport = jurnal.getvalue()
        if avertismente:
            raport += "\n" + "\n".join(avertismente)
        if respinse:
            raport += "\nFisiere ignorate (format nepermis): " + ", ".join(respinse)
        if esuate:
            raport += (
                "\n\nUrmatoarele fisiere NU au putut fi citite si au fost PASTRATE"
                " (nu sterse) pentru verificare manuala:\n"
            )
            for cale in esuate:
                marime = os.path.getsize(cale) if os.path.exists(cale) else 0
                raport += f"   {cale}  ({marime} octeti)\n"

        return render_template("import.html", raport=raport)

    return render_template("import.html", raport=None)


if __name__ == "__main__":
    # host="0.0.0.0" face aplicatia vizibila si de pe alte PC-uri din retea,
    # nu doar de pe acest calculator.
    app.run(host="0.0.0.0", port=5000, debug=True)
