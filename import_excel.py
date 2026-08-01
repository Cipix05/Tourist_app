"""
Script de import: aduce date din unul sau mai multe fisiere Excel (.xls,
.xlsx), Word (.docx, sau .doc vechi - se converteste automat) in baza de
date a aplicatiei. Poti amesteca tipurile de fisiere in aceeasi comanda.

Poti rula scriptul cu mai multe fisiere deodata, chiar daca au coloane
denumite diferit intre ele - cat timp contin informatii similare (nume,
CNP, telefon, adresa pentru turisti; nume, cod, perioada, obiective pentru
excursii), scriptul incearca sa recunoasca automat ce reprezinta fiecare
coloana si fiecare foaie/tabel din fisier.

DEDUPLICARE: daca acelasi turist apare de mai multe ori (in acelasi fisier
sau in fisiere diferite), NU se creeaza un turist nou de fiecare data:
  - daca CNP-ul e identic -> e clar acelasi om, datele se completeaza/combina
  - daca nu exista CNP dar numele e identic -> se presupune acelasi om
  - daca un camp (ex: telefonul) difera fata de ce e deja salvat, cea noua
    valoare se adauga langa cea veche, despartita prin " / ", ca sa nu se
    piarda informatia
  - daca toate datele sunt identice, turistul apare o singura data, fara
    nicio modificare

Utilizare:
    python import_excel.py fisier1.xls
    python import_excel.py fisier1.xls fisier2.doc fisier3.xlsx
"""
import os
import re
import shutil
import subprocess
import sys
import tempfile
import time
import unicodedata
from datetime import datetime

import pandas as pd
from docx import Document as DocxDocument

from app import app, asigura_schema_actualizata
from models import db, Turist, Excursie, Inscriere


# ---------------------------------------------------------------------------
# Recunoasterea coloanelor - fisiere diferite pot numi la fel informatia,
# dar cu alt nume de coloana (ex: "Tel" in loc de "Telefon")
# ---------------------------------------------------------------------------
CANDIDATI_TURIST = {
    "nume": ["nume complet", "nume si prenume", "nume turist", "nume"],
    "cnp": ["cnp"],
    "serie_ci": ["serie ci", "serie si numar", "serie"],
    "emisa_de": ["eliberat de", "emisa de", "emisa"],
    "adresa": ["adresa"],
    "localitate": ["localitate", "oras"],
    "judet": ["judet"],
    "data_nasterii": ["data nasterii", "data nastere", "datanast", "nascut la"],
    "telefon": ["telefon", "nr telefon", "mobil", "tel"],
    "email": ["e-mail", "email", "mail"],
}

CANDIDATI_EXCURSIE = {
    "nume": ["denumire excursie", "denumire", "excursie", "nume"],
    "cod": ["cod"],
    "perioada": ["perioada"],
    "zile": ["nr zile", "zile", "z"],
    "obiective": ["obiective"],
}


def fara_diacritice(text):
    text = unicodedata.normalize("NFKD", str(text))
    return "".join(c for c in text if not unicodedata.combining(c))


def normalizeaza(text):
    return re.sub(r"\s+", " ", fara_diacritice(text)).strip().upper()


def text_curat(val):
    """Converteste orice valoare de celula intr-un string curat sau None."""
    if val is None:
        return None
    if isinstance(val, float):
        if pd.isna(val):
            return None
        if val == int(val):
            return str(int(val))
        return str(val)
    s = str(val).strip()
    return s or None


def combina_valori(veche, noua):
    """Pastreaza valoarea existenta daca cea noua lipseste sau e identica;
    daca difera cu adevarat, le pune pe amandoua despartite prin ' / '."""
    if not noua:
        return veche
    if not veche:
        return noua
    if normalizeaza(veche) == normalizeaza(noua):
        return veche
    parti_existente = {normalizeaza(p) for p in veche.split(" / ")}
    if normalizeaza(noua) in parti_existente:
        return veche
    return f"{veche} / {noua}"


# ---------------------------------------------------------------------------
# Recunoasterea foilor (turisti / excursii / altceva) si a randului de antet
# ---------------------------------------------------------------------------
def gaseste_coloana(coloane, candidati):
    coloane_norm = {normalizeaza(c): c for c in coloane if isinstance(c, str)}
    for cheie in candidati:
        cheie_norm = normalizeaza(cheie)
        for norm, original in coloane_norm.items():
            if norm == cheie_norm or cheie_norm in norm:
                return original
    return None


def mapeaza_coloane(coloane, profil):
    rezultat = {}
    for camp, candidati in profil.items():
        col = gaseste_coloana(coloane, candidati)
        if col:
            rezultat[camp] = col
    return rezultat


def detecteaza_rand_antet(df_brut, max_randuri=6):
    """df_brut citit cu header=None. Cauta randul care pare antet de tabel
    (contine mai multe cuvinte cheie recunoscute)."""
    cuvinte_cheie = set()
    for profil in (CANDIDATI_TURIST, CANDIDATI_EXCURSIE):
        for candidati in profil.values():
            cuvinte_cheie.update(normalizeaza(c) for c in candidati)
    for i in range(min(max_randuri, len(df_brut))):
        valori = [normalizeaza(v) for v in df_brut.iloc[i] if isinstance(v, str)]
        scor = sum(1 for v in valori if any(k in v for k in cuvinte_cheie))
        if scor >= 2:
            return i
    return None


def citeste_foaie(xls, foaie):
    brut = pd.read_excel(xls, sheet_name=foaie, header=None)
    rand_antet = detecteaza_rand_antet(brut)
    if rand_antet is None:
        return pd.read_excel(xls, sheet_name=foaie)
    return pd.read_excel(xls, sheet_name=foaie, header=rand_antet)


def clasifica_foaie(df):
    """Ghiceste tipul foii: 'turisti', 'excursii' sau 'necunoscut'."""
    coloane = list(df.columns)
    coloane_norm = {normalizeaza(c) for c in coloane if isinstance(c, str)}

    # semne clare ca e o foaie de persoane JURIDICE (parteneri, firme de
    # transport etc.) si nu de turisti, chiar daca are si coloane Nume/Adresa
    if any(marcaj in c for c in coloane_norm for marcaj in ("CUI", "IBAN", "ORC")):
        return "necunoscut", {}

    mapare_t = mapeaza_coloane(coloane, CANDIDATI_TURIST)
    mapare_e = mapeaza_coloane(coloane, CANDIDATI_EXCURSIE)

    greutati_t = {"cnp": 3, "adresa": 2, "telefon": 1, "nume": 1}
    greutati_e = {"perioada": 3, "obiective": 2, "cod": 1, "nume": 1}
    scor_t = sum(greutati_t.get(c, 1) for c in mapare_t)
    scor_e = sum(greutati_e.get(c, 1) for c in mapare_e)

    if "nume" in mapare_t and scor_t >= 3 and scor_t >= scor_e:
        return "turisti", mapare_t
    if "nume" in mapare_e and scor_e >= 3 and scor_e > scor_t:
        return "excursii", mapare_e
    return "necunoscut", {}


def valoare(row, mapare, camp):
    col = mapare.get(camp)
    if not col or col not in row:
        return None
    val = row[col]
    return None if pd.isna(val) else val


def parse_data_ro(val):
    if val is None:
        return None
    if not isinstance(val, str):
        try:
            return pd.Timestamp(val).date()
        except Exception:
            return None
    val = val.strip()
    for fmt in ("%d.%m.%Y", "%d-%m-%Y", "%Y-%m-%d"):
        try:
            return datetime.strptime(val, fmt).date()
        except ValueError:
            continue
    return None


def parseaza_perioada(perioada):
    if not isinstance(perioada, str):
        return None, None
    perioada = perioada.strip()
    unica = re.match(r"^(\d{1,2})\.(\d{1,2})\.(\d{4})$", perioada)
    if unica:
        d, m, y = unica.groups()
        try:
            dt = datetime(int(y), int(m), int(d)).date()
            return dt, dt
        except ValueError:
            return None, None
    rang = re.match(r"^(\d{1,2})-(\d{1,2})\.(\d{1,2})\.(\d{4})$", perioada)
    if rang:
        d1, d2, m, y = rang.groups()
        try:
            return datetime(int(y), int(m), int(d1)).date(), datetime(int(y), int(m), int(d2)).date()
        except ValueError:
            return None, None
    return None, None


# ---------------------------------------------------------------------------
# Import turisti, cu deduplicare
# ---------------------------------------------------------------------------
CAMPURI_TURIST_SIMPLE = ["serie_ci", "emisa_de", "adresa", "localitate", "judet", "telefon", "email"]


def importa_turisti(df, mapare, index_nume, index_cnp):
    noi, actualizati = 0, 0
    for _, row in df.iterrows():
        nume = text_curat(valoare(row, mapare, "nume"))
        if not nume:
            continue

        cnp = text_curat(valoare(row, mapare, "cnp"))
        cheie_nume = normalizeaza(nume)

        turist = None
        if cnp and cnp in index_cnp:
            turist = index_cnp[cnp]
        elif cheie_nume in index_nume:
            # daca ambii au CNP si difera, sunt doi oameni diferiti cu acelasi nume
            turist = next(
                (t for t in index_nume[cheie_nume] if not cnp or not t.cnp or t.cnp == cnp), None
            )

        valori_noi = {camp: text_curat(valoare(row, mapare, camp)) for camp in CAMPURI_TURIST_SIMPLE}
        valori_noi["cnp"] = cnp
        data_nasterii = parse_data_ro(valoare(row, mapare, "data_nasterii"))

        if turist is None:
            turist = Turist(nume=nume)
            for camp, val in valori_noi.items():
                setattr(turist, camp, val)
            if data_nasterii:
                turist.data_nasterii = data_nasterii
            db.session.add(turist)
            noi += 1
        else:
            for camp, val in valori_noi.items():
                setattr(turist, camp, combina_valori(getattr(turist, camp), val))
            if data_nasterii and not turist.data_nasterii:
                turist.data_nasterii = data_nasterii
            actualizati += 1

        index_nume.setdefault(cheie_nume, [])
        if turist not in index_nume[cheie_nume]:
            index_nume[cheie_nume].append(turist)
        if turist.cnp:
            index_cnp[turist.cnp] = turist

    db.session.flush()
    print(f"   turisti: {noi} noi, {actualizati} completati (deja existenti, fara duplicare)")


# ---------------------------------------------------------------------------
# Import excursii, cu deduplicare simpla (dupa nume + data de inceput)
# ---------------------------------------------------------------------------
def importa_excursii(df, mapare, index_excursii):
    noi, actualizate = 0, 0
    for _, row in df.iterrows():
        nume = text_curat(valoare(row, mapare, "nume"))
        if not nume:
            continue

        cod = text_curat(valoare(row, mapare, "cod"))
        perioada = valoare(row, mapare, "perioada")
        data_inceput, data_sfarsit = parseaza_perioada(perioada) if perioada else (None, None)
        zile_raw = text_curat(valoare(row, mapare, "zile"))
        obiective = text_curat(valoare(row, mapare, "obiective"))

        cheie_nume = normalizeaza(nume)
        excursie = index_excursii.get((cheie_nume, data_inceput)) or index_excursii.get((cheie_nume, None))

        if excursie is None:
            excursie = Excursie(
                nume=nume, cod=cod, data_inceput=data_inceput, data_sfarsit=data_sfarsit,
                nr_zile=int(zile_raw) if zile_raw and zile_raw.isdigit() else None,
                obiective=obiective,
            )
            db.session.add(excursie)
            noi += 1
        else:
            excursie.cod = excursie.cod or cod
            excursie.obiective = combina_valori(excursie.obiective, obiective)
            actualizate += 1

        index_excursii[(cheie_nume, data_inceput)] = excursie
        index_excursii.setdefault((cheie_nume, None), excursie)

    db.session.flush()
    print(f"   excursii: {noi} noi, {actualizate} completate (deja existente, fara duplicare)")


# ---------------------------------------------------------------------------
# Import inscrieri - foi de tip "Pers.Exc": doar nume turist + nume excursie,
# fara antet formal (formatul vechi, mostenit din baza initiala)
# ---------------------------------------------------------------------------
def importa_inscrieri_foaie_veche(xls, foaie, index_nume, index_excursii):
    df = pd.read_excel(xls, sheet_name=foaie, header=None, skiprows=4,
                        names=["idx", "nume", "excursie", "data", "c5", "c6"])
    adaugate, nepotrivite = 0, []
    perechi_facute = set()

    for _, row in df.iterrows():
        nume, excursie_nume = row["nume"], row["excursie"]
        if not isinstance(nume, str) or not isinstance(excursie_nume, str):
            continue

        candidati = index_nume.get(normalizeaza(nume), [])
        turist = candidati[0] if candidati else None
        excursie = index_excursii.get((normalizeaza(excursie_nume), None))

        if not turist or not excursie:
            nepotrivite.append((nume, excursie_nume))
            continue

        db.session.flush()
        exista = (
            Inscriere.query.filter_by(turist_id=turist.id, excursie_id=excursie.id).first()
            if turist.id and excursie.id else None
        )
        cheie = (turist.id, excursie.id)
        if exista or cheie in perechi_facute:
            continue
        perechi_facute.add(cheie)

        db.session.add(Inscriere(turist=turist, excursie=excursie, suma_achitata=0))
        adaugate += 1

    db.session.flush()
    print(f"   inscrieri (din '{foaie}'): {adaugate} adaugate")
    if nepotrivite:
        print(f"   ATENTIE: {len(nepotrivite)} inscrieri nu s-au putut asocia automat (nume care nu s-a gasit exact):")
        for nume, exc in nepotrivite[:20]:
            print(f"      - {nume}  ->  {exc}")
        if len(nepotrivite) > 20:
            print(f"      ... si inca {len(nepotrivite) - 20}")


# ---------------------------------------------------------------------------
# Fisiere .doc vechi (format binar, dinainte de Word 2007) - trebuie
# convertite in .docx inainte sa le putem citi. Incercam doua metode,
# in ordine, in functie de ce e instalat pe calculator.
# ---------------------------------------------------------------------------
def converteste_doc_in_docx(cale_doc):
    """Incearca sa converteasca un .doc in .docx, folosind (in ordine):
    Microsoft Word (daca e instalat, prin pywin32) sau LibreOffice (daca
    e instalat). Returneaza calea catre fisierul .docx temporar creat,
    sau None daca nicio metoda n-a functionat."""
    cale_doc = os.path.abspath(cale_doc)
    folder_temp = tempfile.mkdtemp(prefix="import_doc_")

    # Metoda 1: Microsoft Word, prin automatizare COM (doar Windows, doar
    # daca ai Word instalat - e metoda cea mai sigura pe un PC obisnuit)
    try:
        import win32com.client
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        try:
            doc = word.Documents.Open(cale_doc, ReadOnly=True)
            cale_docx = os.path.join(folder_temp, "convertit.docx")
            doc.SaveAs(cale_docx, FileFormat=16)  # 16 = wdFormatXMLDocument (.docx)
            doc.Close(False)
            if os.path.exists(cale_docx):
                return cale_docx
        finally:
            word.Quit()
    except Exception:
        pass  # Word nu e instalat sau a esuat - incercam urmatoarea metoda

    # Metoda 2: LibreOffice, din linia de comanda (daca e instalat)
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice:
        try:
            subprocess.run(
                [soffice, "--headless", "--convert-to", "docx", "--outdir", folder_temp, cale_doc],
                check=True, timeout=180, capture_output=True,
            )
            nume_baza = os.path.splitext(os.path.basename(cale_doc))[0]
            cale_docx = os.path.join(folder_temp, nume_baza + ".docx")
            if os.path.exists(cale_docx):
                return cale_docx
        except Exception:
            pass

    return None


# ---------------------------------------------------------------------------
# Citire fisiere Word (.docx) - fiecare tabel din document e tratat ca o
# "foaie" separata, exact ca la Excel
# ---------------------------------------------------------------------------
# ---------------------------------------------------------------------------
# Citire fisiere Word (.docx) - atat tabelele (fiecare tratat ca o "foaie"
# separata, exact ca la Excel), cat si paragrafele din afara tabelelor
# (necesare pentru formatul "o excursie + participantii ei", care are
# denumirea si perioada scrise ca text liber deasupra tabelului)
# ---------------------------------------------------------------------------
def citeste_docx_complet(cale):
    """Returneaza (paragrafe, tabele). 'tabele' e o lista de
    (nume_tabel, DataFrame), la fel ca inainte.

    Incearca de 3 ori, cu o mica pauza intre incercari - pe un folder de
    retea (NAS/share), uneori fisierul e "salvat" dar nu e inca vizibil
    complet cand incepe citirea imediat dupa."""
    ultima_eroare = None
    for incercare in range(3):
        if incercare > 0:
            time.sleep(0.5)
        try:
            doc = DocxDocument(cale)
            break
        except Exception as exc:
            ultima_eroare = exc
    else:
        raise ultima_eroare

    paragrafe = [p.text for p in doc.paragraphs]

    tabele = []
    for idx, tabel in enumerate(doc.tables, start=1):
        randuri = [[celula.text.strip() for celula in rand.cells] for rand in tabel.rows]
        randuri = [r for r in randuri if any(v for v in r)]  # sare peste randuri complet goale
        if len(randuri) < 2:
            continue  # doar antet sau tabel gol, nimic de importat
        antet, *date = randuri
        df = pd.DataFrame(date, columns=antet).replace("", None)
        tabele.append((f"Tabel {idx}", df))

    return paragrafe, tabele


# ---------------------------------------------------------------------------
# Format special 1: document cu "turisti premiati" - o coloana cu numele
# turistului, o coloana EXACT numita "N" (total excursii cu agentia) si o
# coloana care contine excursiile primite cadou (text liber, o linie per
# excursie). NU creeaza turisti noi - e o corectie pentru cei deja existenti,
# seteaza ajustare_excursii/ajustare_cadouri ca totalul din aplicatie sa
# coincida cu totalul din document.
# ---------------------------------------------------------------------------
def detecteaza_tabel_corectii(df):
    coloane = [c for c in df.columns if isinstance(c, str)]
    coloane_norm = {normalizeaza(c): c for c in coloane}

    # "N" trebuie sa se potriveasca EXACT, nu ca substring (altfel s-ar
    # potrivi din greseala cu orice coloana care contine litera N, ex "Nume")
    col_total = coloane_norm.get("N")
    if not col_total:
        return None

    col_nume = gaseste_coloana(coloane, CANDIDATI_TURIST["nume"])
    if not col_nume:
        return None

    col_cadou = None
    for norm, orig in coloane_norm.items():
        if "CADOU" in norm:
            col_cadou = orig
            break
    if not col_cadou:
        return None

    return {"nume": col_nume, "total": col_total, "cadou": col_cadou}


def importa_corectii(df, mapare, index_nume):
    actualizati, negasiti = 0, []
    for _, row in df.iterrows():
        nume = text_curat(row.get(mapare["nume"]))
        if not nume:
            continue

        candidati = index_nume.get(normalizeaza(nume), [])
        turist = candidati[0] if candidati else None
        if not turist:
            negasiti.append(nume)
            continue

        total_text = text_curat(row.get(mapare["total"]))
        total = int(total_text) if total_text and total_text.isdigit() else None

        cadou_text = row.get(mapare["cadou"])
        cadou_text = "" if pd.isna(cadou_text) else str(cadou_text).strip()
        nr_cadou = len([linie for linie in cadou_text.split("\n") if linie.strip()]) if cadou_text else 0

        # "N" (total) din document INCLUDE deja excursiile cadou - trebuie
        # scazute de-acolo inainte sa calculam partea "platita", altfel
        # cadourile s-ar numara de doua ori in total afisat de aplicatie.
        deja_total = len(turist.inscrieri)
        deja_cadou = len([i for i in turist.inscrieri if i.excursie_cadou])
        deja_platite = deja_total - deja_cadou

        if total is not None:
            total_platite_document = total - nr_cadou
            turist.ajustare_excursii = total_platite_document - deja_platite

        turist.ajustare_cadouri = nr_cadou - deja_cadou

        actualizati += 1

    print(f"   corectii aplicate: {actualizati} turisti actualizati (numar total excursii + cadouri)")
    if negasiti:
        print(f"   ATENTIE: {len(negasiti)} nume din document nu s-au gasit in baza de date (turist inexistent - trebuie adaugat separat, apoi rulezi din nou acest fisier):")
        for n in negasiti[:20]:
            print(f"      - {n}")
        if len(negasiti) > 20:
            print(f"      ... si inca {len(negasiti) - 20}")


# ---------------------------------------------------------------------------
# Format special 2: document cu O SINGURA excursie si toti participantii ei
# (denumirea + perioada scrise ca text liber deasupra tabelului, de forma
# "<denumire> din perioada DD-DD.MM.YYYY" sau "<denumire> din perioada
# DD.MM-DD.MM.YYYY"). Tabelul contine Nume, Telefon, E-mail, Achitat.
# ---------------------------------------------------------------------------
def extrage_nume_perioada_participanti(paragrafe):
    """Cauta tiparul '<denumire> din perioada <date>' in paragrafele
    dinaintea tabelului. Returneaza (nume, data_inceput, data_sfarsit) -
    toate None daca tiparul nu a fost gasit (document de alt fel)."""
    text_complet = "\n".join(p for p in paragrafe if p.strip())

    # perioada in aceeasi luna: "29-30.07.2015"
    m = re.search(
        r"(.+?)\s+din\s+perioada\s+(\d{1,2})\s*-\s*(\d{1,2})\.(\d{1,2})\.(\d{4})",
        text_complet, re.IGNORECASE,
    )
    if m:
        nume, d1, d2, luna, an = m.groups()
        try:
            return (
                nume.strip(" :-\n"),
                datetime(int(an), int(luna), int(d1)).date(),
                datetime(int(an), int(luna), int(d2)).date(),
            )
        except ValueError:
            pass

    # perioada intre luni diferite: "30.07-02.08.2015" sau "30.07.2015-02.08.2015"
    m = re.search(
        r"(.+?)\s+din\s+perioada\s+(\d{1,2})\.(\d{1,2})\.?(\d{4})?\s*-\s*(\d{1,2})\.(\d{1,2})\.(\d{4})",
        text_complet, re.IGNORECASE,
    )
    if m:
        nume, d1, m1, an1, d2, m2, an2 = m.groups()
        an1 = an1 or an2
        try:
            return (
                nume.strip(" :-\n"),
                datetime(int(an1), int(m1), int(d1)).date(),
                datetime(int(an2), int(m2), int(d2)).date(),
            )
        except ValueError:
            pass

    return None, None, None


def detecteaza_tabel_participanti(df):
    """Coloana 'Achitat' e semnul distinctiv al acestui format - fara ea,
    nu incercam sa tratam tabelul ca fiind de acest tip (ca sa nu confundam
    cu o lista generica de turisti, care are aceleasi coloane nume/telefon)."""
    coloane = [c for c in df.columns if isinstance(c, str)]
    col_achitat = gaseste_coloana(coloane, ["achitat", "suma achitata", "plata"])
    col_nume = gaseste_coloana(coloane, CANDIDATI_TURIST["nume"])
    if not col_achitat or not col_nume:
        return None
    return {
        "nume": col_nume,
        "telefon": gaseste_coloana(coloane, CANDIDATI_TURIST["telefon"]),
        "email": gaseste_coloana(coloane, CANDIDATI_TURIST["email"]),
        "achitat": col_achitat,
    }


def importa_excursie_cu_participanti(nume, data_inceput, data_sfarsit, df, mapare,
                                      index_excursii, index_nume, index_cnp):
    cheie_nume = normalizeaza(nume)
    excursie = index_excursii.get((cheie_nume, data_inceput)) or index_excursii.get((cheie_nume, None))

    if excursie is None:
        excursie = Excursie(nume=nume, data_inceput=data_inceput, data_sfarsit=data_sfarsit)
        db.session.add(excursie)
        db.session.flush()
        index_excursii[(cheie_nume, data_inceput)] = excursie
        index_excursii.setdefault((cheie_nume, None), excursie)
        info_data = f" ({data_inceput.strftime('%d.%m.%Y')})" if data_inceput else " (perioada nerecunoscuta - adaug-o manual din aplicatie)"
        print(f"   Excursie noua: '{nume}'{info_data}")
    else:
        print(f"   Excursie deja existenta: '{excursie.nume}' - adaug participantii la ea")

    adaugati, actualizati = 0, 0
    for _, row in df.iterrows():
        nume_turist = text_curat(row.get(mapare["nume"]))
        if not nume_turist:
            continue

        telefon = text_curat(row.get(mapare["telefon"])) if mapare.get("telefon") else None
        email_brut = text_curat(row.get(mapare["email"])) if mapare.get("email") else None
        # unele fisiere au din greseala un al doilea numar de telefon in
        # coloana de e-mail - il pastram, doar nu il punem in campul email
        email = email_brut if email_brut and "@" in email_brut else None
        telefon_secundar = email_brut if email_brut and "@" not in email_brut else None

        cheie_t = normalizeaza(nume_turist)
        candidati = index_nume.get(cheie_t, [])
        turist = candidati[0] if candidati else None

        if turist is None:
            turist = Turist(nume=nume_turist, telefon=telefon, email=email)
            if telefon_secundar:
                turist.telefon = combina_valori(turist.telefon, telefon_secundar)
            db.session.add(turist)
            db.session.flush()
            index_nume.setdefault(cheie_t, []).append(turist)
            if turist.cnp:
                index_cnp[turist.cnp] = turist
        else:
            turist.telefon = combina_valori(turist.telefon, telefon)
            if telefon_secundar:
                turist.telefon = combina_valori(turist.telefon, telefon_secundar)
            turist.email = combina_valori(turist.email, email)
            actualizati += 1

        achitat_text = text_curat(row.get(mapare["achitat"]))
        suma = None
        if achitat_text:
            try:
                suma = float(achitat_text.replace(",", "."))
            except ValueError:
                suma = None

        db.session.flush()
        exista = Inscriere.query.filter_by(turist_id=turist.id, excursie_id=excursie.id).first()
        if exista:
            if suma is not None and not exista.suma_achitata:
                exista.suma_achitata = suma
            continue

        db.session.add(Inscriere(turist=turist, excursie=excursie, suma_achitata=suma or 0))
        adaugati += 1

        # aceasta excursie e (aproape sigur) deja numarata in totalul din
        # documentul de "turisti premiati" - ca sa nu se numere de doua
        # ori, scadem 1 din corectia manuala de fiecare data cand o
        # inscriere REALA (cu excursie si data concreta) o inlocuieste pe
        # cea "generica", inregistrata doar ca numar in corectie.
        if (turist.ajustare_excursii or 0) > 0:
            turist.ajustare_excursii -= 1

    db.session.flush()
    print(f"   participanti: {adaugati} inscrieri noi, {actualizati} turisti completati (deja existenti)")


def proceseaza_fisier(cale, index_nume, index_cnp, index_excursii, foi_inscrieri_amanate):
    print(f"\n=== {cale} ===")
    extensie = cale.lower().rsplit(".", 1)[-1] if "." in cale else ""

    if extensie == "doc":
        print("   Fisier .doc (format vechi) - incerc sa-l convertesc in .docx...")
        cale_convertita = converteste_doc_in_docx(cale)
        if not cale_convertita:
            print(
                "   Nu am putut converti automat acest fisier (nu am gasit nici Word,\n"
                "   nici LibreOffice instalate). Solutie rapida: deschide fisierul in\n"
                "   Word, foloseste 'Save As' -> 'Word Document (*.docx)', apoi\n"
                "   ruleaza din nou importul cu noul fisier .docx."
            )
            return
        print("   Convertit cu succes, continui importul.")
        extensie = "docx"
        cale = cale_convertita

    if extensie == "docx":
        try:
            paragrafe, tabele = citeste_docx_complet(cale)
        except Exception as exc:
            print(f"   Nu am putut deschide fisierul Word: {exc}")
            return
        if not tabele:
            print("   Nu am gasit niciun tabel folosibil in acest document Word.")
            return

        # 1) formatul "o excursie + participantii ei" (are neaparat coloana
        # 'Achitat' SI un tipar '<nume> din perioada <date>' in text) - cel
        # mai specific dintre toate, il verificam primul
        nume_exc, data_i, data_f = extrage_nume_perioada_participanti(paragrafe)
        if nume_exc:
            for nume_tabel, df in tabele:
                mapare_part = detecteaza_tabel_participanti(df)
                if mapare_part:
                    print(f"   Recunoscut ca fisier de participanti la o singura excursie.")
                    importa_excursie_cu_participanti(
                        nume_exc, data_i, data_f, df, mapare_part,
                        index_excursii, index_nume, index_cnp,
                    )
                    return

        # 2) formatul "turisti premiati" (corectii de numar excursii/cadouri)
        for nume_tabel, df in tabele:
            mapare_corectii = detecteaza_tabel_corectii(df)
            if mapare_corectii:
                print(f"   {nume_tabel} recunoscut ca lista de corectii (nume + total excursii + cadouri).")
                importa_corectii(df, mapare_corectii, index_nume)
                return

        # 3) fallback: tabele generice de turisti/excursii, ca inainte
        for nume_tabel, df in tabele:
            tip, mapare = clasifica_foaie(df)
            if tip == "turisti":
                print(f"   {nume_tabel} recunoscut ca turisti.")
                importa_turisti(df, mapare, index_nume, index_cnp)
            elif tip == "excursii":
                print(f"   {nume_tabel} recunoscut ca excursii.")
                importa_excursii(df, mapare, index_excursii)
            else:
                print(f"   {nume_tabel}: nu am recunoscut formatul, il ignor.")
        return

    if extensie not in ("xls", "xlsx"):
        print(f"   Format de fisier necunoscut ('.{extensie}') - accept .xls, .xlsx, .doc sau .docx.")
        return

    try:
        xls = pd.ExcelFile(cale)
    except Exception as exc:
        print(f"   Nu am putut deschide fisierul: {exc}")
        return

    for foaie in xls.sheet_names:
        nume_foaie_norm = normalizeaza(foaie)
        if "EXC" in nume_foaie_norm and "PERS" in nume_foaie_norm:
            # foaie de tip "Pers.Exc" (inscrieri) - se proceseaza dupa ce
            # avem toti turistii si excursiile din toate fisierele
            foi_inscrieri_amanate.append((cale, foaie))
            continue
        try:
            df = citeste_foaie(xls, foaie)
        except Exception as exc:
            print(f"   Foaia '{foaie}': nu am putut-o citi ({exc}), o ignor.")
            continue

        tip, mapare = clasifica_foaie(df)
        if tip == "turisti":
            print(f"   Foaia '{foaie}' recunoscuta ca turisti.")
            importa_turisti(df, mapare, index_nume, index_cnp)
        elif tip == "excursii":
            print(f"   Foaia '{foaie}' recunoscuta ca excursii.")
            importa_excursii(df, mapare, index_excursii)
        else:
            print(f"   Foaia '{foaie}': nu am recunoscut formatul, o ignor.")


def importa_fisiere(cai_fisiere):
    """Logica de import reutilizabila - apelata atat din linia de comanda
    (main(), mai jos), cat si din aplicatia web (ruta /import), fara sa
    duplice codul."""
    with app.app_context():
        asigura_schema_actualizata()

        # index-uri construite din ce e deja in baza de date, ca sa
        # deduplicam si fata de importuri anterioare, nu doar in cadrul
        # acestei rulari
        index_nume = {}
        for t in Turist.query.all():
            index_nume.setdefault(normalizeaza(t.nume), []).append(t)
        index_cnp = {t.cnp: t for t in Turist.query.all() if t.cnp}
        index_excursii = {}
        for e in Excursie.query.all():
            index_excursii[(normalizeaza(e.nume), e.data_inceput)] = e
            index_excursii.setdefault((normalizeaza(e.nume), None), e)

        foi_inscrieri_amanate = []
        for fisier in cai_fisiere:
            proceseaza_fisier(fisier, index_nume, index_cnp, index_excursii, foi_inscrieri_amanate)

        db.session.commit()

        for cale, foaie in foi_inscrieri_amanate:
            print(f"\n=== inscrieri din {cale} / '{foaie}' ===")
            xls = pd.ExcelFile(cale)
            importa_inscrieri_foaie_veche(xls, foaie, index_nume, index_excursii)
        db.session.commit()

    print("\nImport finalizat.")


def main():
    if len(sys.argv) < 2:
        print("Utilizare: python import_excel.py fisier1.xls [fisier2.xlsx ...]")
        sys.exit(1)
    importa_fisiere(sys.argv[1:])


if __name__ == "__main__":
    main()
